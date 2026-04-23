"""
Tạo file Word báo giá L&A E-Learning cho sếp.
2 bảng: Input/Output + Bảng giá phân cấp chỉnh sửa.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Cấu hình style ──
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(10)

NAVY = RGBColor(0x1A, 0x2B, 0x50)      # Xanh navy đậm
GOLD = RGBColor(0xD4, 0xA8, 0x1F)      # Vàng gold
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_BG = RGBColor(0xF5, 0xF5, 0xF5)

def set_cell_shading(cell, hex_color):
    """Đặt màu nền cho ô."""
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    )

def set_cell_border(cell, **kwargs):
    """Đặt border cho ô."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_run(paragraph, text, bold=False, size=10, color=BLACK, font_name="Arial"):
    """Thêm text có format vào paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def format_header_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Format ô header với nền navy, chữ trắng."""
    set_cell_shading(cell, "1A2B50")
    p = cell.paragraphs[0]
    p.alignment = align
    add_formatted_run(p, text, bold=True, size=11, color=WHITE)

# ============================================================
# TIÊU ĐỀ TÀI LIỆU
# ============================================================
title = doc.add_heading("BÁO GIÁ DỰ ÁN", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = NAVY
    run.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(subtitle, "L&A E-LEARNING PLATFORM", bold=True, size=14, color=NAVY)
doc.add_paragraph()  # Khoảng trống

# Thông tin dự án
info = doc.add_paragraph()
add_formatted_run(info, "Dự án: ", bold=True, size=11)
add_formatted_run(info, "L&A E-Learning Onboarding Platform", size=11)
info2 = doc.add_paragraph()
add_formatted_run(info2, "Nền tảng: ", bold=True, size=11)
add_formatted_run(info2, "Open edX (Tutor) + React Frontend", size=11)
info3 = doc.add_paragraph()
add_formatted_run(info3, "Ngày: ", bold=True, size=11)
add_formatted_run(info3, "22/04/2026", size=11)

doc.add_paragraph()  # Khoảng trống

# ============================================================
# BẢNG 1: INPUT / OUTPUT
# ============================================================
h1 = doc.add_heading("1. BẢNG THÀNH PHẨM BÀN GIAO", level=1)
for run in h1.runs:
    run.font.color.rgb = NAVY

table1 = doc.add_table(rows=3, cols=2)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.autofit = True

# Đặt chiều rộng cột
for row in table1.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(12)

# ── Header ──
format_header_cell(table1.cell(0, 0), "L&A Input")
format_header_cell(table1.cell(0, 1), "Nesso Output\n(Thành phẩm bàn giao)")

# ── Row 1: Content Visualization ──
cell_input = table1.cell(1, 0)
set_cell_shading(cell_input, "F5F5F5")
p_in = cell_input.paragraphs[0]
add_formatted_run(p_in, "34 slides PPTX tĩnh", bold=True, size=10)

cell_output1 = table1.cell(1, 1)
p1 = cell_output1.paragraphs[0]
add_formatted_run(p1, "Trọn gói E-LEARNING CONTENT VISUALIZATION", bold=True, size=11)
add_formatted_run(p1, ", bao gồm:", size=10)

# Bullet points
items_viz = [
    ("Content Setup: ", "Hệ thống hóa nội dung logic, mạch lạc."),
    ("Visual Design: ", "12 trang Graphic & 10 trang Infographic sống động, dễ nhớ."),
    ("Interactive Web: ", "Quiz Test sau mỗi module, Game, Streak (Tăng ghi nhớ)."),
    ("Multimedia: ", "05 Videos 2D Motion & AI Voiceover chuyên nghiệp."),
]
for label, desc in items_viz:
    bp = cell_output1.add_paragraph(style="List Bullet")
    add_formatted_run(bp, label, bold=True, size=9)
    add_formatted_run(bp, desc, size=9)

# ── Row 2: E-Learning System ──
cell_input2 = table1.cell(2, 0)
set_cell_shading(cell_input2, "F5F5F5")
# Merge với cell trên (để trống)

cell_output2 = table1.cell(2, 1)
p2 = cell_output2.paragraphs[0]
add_formatted_run(p2, "Trọn gói E-LEARNING SYSTEM", bold=True, size=11)
add_formatted_run(p2, ", bao gồm:", size=10)

items_sys = [
    "Nền tảng Open edX (LMS + CMS Studio) triển khai trên Cloud",
    "Frontend React tùy chỉnh (Dashboard, Course Viewer, Quiz Engine)",
    "OAuth2 Authentication + Role-based Routing (Learner / Mentor / Admin)",
    "Token Auto-Refresh + Encrypted Storage (Bảo mật cấp doanh nghiệp)",
    "XBlock Content Rendering: Video Playback, HTML Content, Interactive Quiz",
    "Progress Tracking & Completion System (API thật, không mock data)",
    "Responsive Design Dark/Light mode",
    "ĐỘI FREELANCER CUNG CẤP",
]
for item in items_sys:
    bp = cell_output2.add_paragraph(style="List Bullet")
    if item == "ĐỘI FREELANCER CUNG CẤP":
        add_formatted_run(bp, item, bold=True, size=9)
    else:
        add_formatted_run(bp, item, size=9)

doc.add_paragraph()  # Khoảng trống

# ============================================================
# BẢNG 2: BẢNG GIÁ PHÂN CẤP CHỈNH SỬA
# ============================================================
h2 = doc.add_heading("2. BẢNG GIÁ PHÂN CẤP CHỈNH SỬA", level=1)
for run in h2.runs:
    run.font.color.rgb = NAVY

# Tạo bảng header riêng (banner vàng)
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner.cell(0, 0)
set_cell_shading(banner_cell, "D4A81F")
bp = banner_cell.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(bp, "BẢNG GIÁ PHÂN CẤP CHỈNH SỬA", bold=True, size=14, color=WHITE)

# Tạo bảng chính
table2 = doc.add_table(rows=1, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = True

# Header
headers = ["Cấp độ", "Tên gọi", "Chi tiết kỹ thuật (Phạm vi công việc)"]
for i, h in enumerate(headers):
    cell = table2.cell(0, i)
    set_cell_shading(cell, "1A2B50")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, h, bold=True, size=10, color=WHITE)

# ── LEVEL 1 — BASIC ──
row = table2.add_row()
row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row.cells[0].paragraphs[0], "LEVEL 1", bold=True, size=10, color=NAVY)
row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row.cells[1].paragraphs[0], "BASIC", bold=True, size=11, color=NAVY)
set_cell_shading(row.cells[1], "F0F4FF")

level1_items = [
    ("UI/UX: ", "Sửa lỗi hiển thị, chỉnh alignment, spacing, typo trên giao diện."),
    ("Config: ", "Cập nhật biến môi trường (.env), thay đổi cấu hình proxy, timeout."),
    ("Content: ", "Thay đổi text/label/placeholder, cập nhật hình ảnh tĩnh (không đổi layout)."),
    ("Style: ", "Cập nhật màu sắc, font, theme (Dark/Light) theo Brand Guidelines."),
    ("Hotfix: ", "Sửa bug nhỏ không ảnh hưởng logic nghiệp vụ (CSS, responsive)."),
]
for label, desc in level1_items:
    bp = row.cells[2].add_paragraph(style="List Bullet")
    add_formatted_run(bp, label, bold=True, size=9)
    add_formatted_run(bp, desc, size=9)

# ── LEVEL 2 — MEDIUM ──
row2 = table2.add_row()
row2.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row2.cells[0].paragraphs[0], "LEVEL 2", bold=True, size=10, color=NAVY)
row2.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row2.cells[1].paragraphs[0], "MEDIUM", bold=True, size=11, color=RGBColor(0xD4, 0xA8, 0x1F))
set_cell_shading(row2.cells[1], "FFF8E7")

level2_items = [
    ("Frontend: ", "Thêm/sửa React component, thay đổi logic hiển thị (dưới 30% code trong module)."),
    ("API: ", "Tích hợp thêm API endpoint mới từ Open edX (courses, enrollment, progress)."),
    ("Auth: ", "Sửa logic authentication, role-based routing, token refresh flow."),
    ("Feature: ", "Thêm tính năng nhỏ: bộ lọc, sắp xếp, phân trang, search khóa học."),
    ("Bug Backend: ", "Fix lỗi liên quan API response, CORS, data transformation (dưới 30% logic)."),
]
for label, desc in level2_items:
    bp = row2.cells[2].add_paragraph(style="List Bullet")
    add_formatted_run(bp, label, bold=True, size=9)
    add_formatted_run(bp, desc, size=9)

# ── LEVEL 3 — COMPLEX ──
row3 = table2.add_row()
row3.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row3.cells[0].paragraphs[0], "LEVEL 3", bold=True, size=10, color=NAVY)
row3.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(row3.cells[1].paragraphs[0], "COMPLEX", bold=True, size=11, color=RGBColor(0xE0, 0x40, 0x40))
set_cell_shading(row3.cells[1], "FFF0F0")

level3_items = [
    ("Architecture: ", "Thay đổi kiến trúc hệ thống, refactor module lớn (>50% codebase), migration framework."),
    ("Integration: ", "Tích hợp hệ thống bên thứ 3 mới (SSO, Payment Gateway, LRS/xAPI, Notification Service)."),
    ("Infrastructure: ", "Thay đổi hạ tầng triển khai: Docker/K8s config, CI/CD pipeline, domain/SSL, scale server."),
    ("Security: ", "Overhaul bảo mật: thay đổi cơ chế mã hóa token, audit log, CSP headers, penetration fix."),
    ("Database: ", "Thay đổi schema Open edX, custom plugin Tutor, migration dữ liệu giữa các môi trường."),
    ("Full Rebuild: ", "Xây lại toàn bộ module (Quiz Engine, Video Player, Dashboard) với yêu cầu nghiệp vụ mới."),
]
for label, desc in level3_items:
    bp = row3.cells[2].add_paragraph(style="List Bullet")
    add_formatted_run(bp, label, bold=True, size=9)
    add_formatted_run(bp, desc, size=9)

# Đặt chiều rộng cột cho bảng 2
for row in table2.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(3.5)
    row.cells[2].width = Cm(11)

# ============================================================
# LƯU FILE
# ============================================================
output_path = r"d:\A-LA-DEMO\BAO_GIA_LA_ELEARNING_v2.docx"
doc.save(output_path)
print(f"[OK] Da tao file: {output_path}")
print(f"    Kich thuoc: {os.path.getsize(output_path):,} bytes")
